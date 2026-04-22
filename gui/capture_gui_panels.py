#!/usr/bin/env python3
"""Capture GUI-level screenshots for TC01..TC17 from the IDCE web app.

This script drives the browser, runs each testcase in the editor, and captures
four output panels directly from the GUI:
- dashboard
- code compare
- cfg view
- analysis logs

Optional: generate a docx/md report from the captured screenshots.

Requirements:
  pip install playwright python-docx
  playwright install chromium

Example:
  python3 capture_gui_panels.py \
    --url http://localhost:3000 \
    --out-dir /media/vyrion/Data1/Nitin/Documents/tmp3 \
    --report /media/vyrion/Data1/Nitin/Documents/tmp3/tc_report.docx
"""

from __future__ import annotations

import argparse
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple

TESTCASES: Dict[str, str] = {
    "TC01": """#include <iostream>
using namespace std;

int main() {
    int x = 5;
    int y = x + 2;
    cout << x << endl;
    return 0;
}""",
    "TC02": """#include <iostream>
using namespace std;

int main() {
    if (0) {
        int x = 5;
    }
    cout << \"Done\";
    return 0;
}""",
    "TC03": """#include <iostream>
using namespace std;

int main() {
    int a = 2;
    int b = a + 3;
    int c = b * 4;
    cout << a << endl;
    return 0;
}""",
    "TC04": """#include <iostream>
using namespace std;

int main() {
    int a = 3;
    int b = a + 2;
    int c = b * 2;
    cout << b;
    return 0;
}""",
    "TC05": """#include <iostream>
using namespace std;

int main() {
    int sum = 0;
    for (int i = 0; i < 5; i++) {
        int temp = i * 3;
        sum += i;
    }
    cout << sum;
    return 0;
}""",
    "TC06": """#include <iostream>
using namespace std;

int main() {
    int x = 10;
    cout << \"Hello\";
    int y = x + 5;
    return 0;
}""",
    "TC07": """#include <iostream>
using namespace std;

int main() {
    int a = 5;
    int *p = &a;
    *p = 20;
    return 0;
}""",
    "TC08": """#include <iostream>
using namespace std;

int main() {
    int x = 0;
    if (x) {
        cout << \"A\";
    } else {
        cout << \"B\";
    }
    return 0;
}""",
    "TC09": """#include <iostream>
using namespace std;

void f() {
    cout << \"Side\";
}

int main() {
    f();
    int x = 5;
    return 0;
}""",
    "TC10": """#include <iostream>
using namespace std;

int main() {
    int a = 4;
    int b = a + 3;
    int c = b + 2;
    cout << a;
    return 0;
}""",
    "TC11": """#include <iostream>
using namespace std;

int main() {
    if (1) {
        if (0) {
            int x = 5;
        }
    }
    return 0;
}""",
    "TC12": """#include <iostream>
using namespace std;

int main() {
    int a = 1;
    int b = 2;
    int c = a + b;
    int d = c + 5;
    cout << a;
    return 0;
}""",
    "TC13": """#include <iostream>
using namespace std;

int main() {
    int sum = 0;
    for (int i = 0; i < 5; i++) {
        int x = i * 2;
        sum += x;
        int y = i + 10;
    }
    cout << sum;
    return 0;
}""",
    "TC14": """#include <iostream>
using namespace std;

int f(int x) {
    return x * 2;
}

int main() {
    int val = f(5);
    return 0;
}""",
    "TC15": """#include <iostream>
using namespace std;

int main() {
    int x = 5;
    return 0;
    int y = x + 2;
}""",
    "TC16": """#include <iostream>
using namespace std;

int f(int n) {
    if (n == 0) return 0;
    return f(n - 1);
}

int main() {
    f(3);
    return 0;
}""",
    "TC17": """#include <iostream>
using namespace std;

int main() {
    int a = 5;
    int b = a + 2;
    if (0) {
        int c = b * 2;
    }
    cout << a;
    return 0;
}""",
}

PANEL_MAP: List[Tuple[str, str, str]] = [
    ("dashboard", "#panel-dashboard", "dashboard"),
    ("compare", "#panel-compare", "code-compare"),
    ("cfg", "#panel-cfg", "cfg-view"),
    ("logs", "#panel-logs", "analysis-logs"),
]

PANEL_ORDER = [
    ("dashboard", "Dashboard"),
    ("code-compare", "Code Compare"),
    ("cfg-view", "CFG View"),
    ("analysis-logs", "Analysis Logs"),
]

PANEL_ALIASES = {
    "dashboard": {"dashboard", "p1", "panel1"},
    "code-compare": {"code-compare", "code_compare", "compare", "p2", "panel2"},
    "cfg-view": {"cfg-view", "cfg_view", "cfg", "p3", "panel3"},
    "analysis-logs": {"analysis-logs", "analysis_logs", "logs", "p4", "panel4"},
}


def tc_keys(start: int, end: int) -> List[str]:
    return [f"TC{i:02d}" for i in range(start, end + 1)]


def normalize_panel(raw_suffix: str) -> Optional[str]:
    suffix = raw_suffix.strip().lower()
    for canonical, aliases in PANEL_ALIASES.items():
        if suffix in aliases:
            return canonical
    return None


def collect_images(images_dir: Path) -> Dict[str, Dict[str, Path]]:
    by_tc: Dict[str, Dict[str, Path]] = {}
    for p in sorted(images_dir.iterdir()):
        if not p.is_file() or p.suffix.lower() not in {".png", ".jpg", ".jpeg", ".webp"}:
            continue
        stem = p.stem
        if "_" not in stem:
            continue
        tc, suffix = stem.split("_", 1)
        tc = tc.upper()
        panel = normalize_panel(suffix)
        if not tc.startswith("TC") or panel is None:
            continue
        by_tc.setdefault(tc, {})[panel] = p
    return by_tc


def render_markdown(output: Path, title: str, tcs: List[str], images: Dict[str, Dict[str, Path]]) -> None:
    lines: List[str] = [f"# {title}", ""]
    for tc in tcs:
        lines.append(f"## {tc}")
        lines.append("")
        tc_imgs = images.get(tc, {})
        for panel_key, panel_label in PANEL_ORDER:
            lines.append(f"### {panel_label}")
            img = tc_imgs.get(panel_key)
            if img:
                lines.append(f"![{tc} {panel_label}]({img.as_posix()})")
            else:
                lines.append("_Missing screenshot_")
            lines.append("")
    output.write_text("\n".join(lines), encoding="utf-8")


def render_docx(output: Path, title: str, tcs: List[str], images: Dict[str, Dict[str, Path]], width_inches: float = 6.0) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as e:
        raise RuntimeError("python-docx is required for .docx output. Install with: pip install python-docx") from e

    doc = Document()
    doc.add_heading(title, level=0)

    for idx, tc in enumerate(tcs):
        doc.add_heading(tc, level=1)
        tc_imgs = images.get(tc, {})
        for panel_key, panel_label in PANEL_ORDER:
            doc.add_paragraph(panel_label)
            img = tc_imgs.get(panel_key)
            if img and img.exists():
                doc.add_picture(str(img), width=Inches(width_inches))
            else:
                doc.add_paragraph("Missing screenshot")
        if idx < len(tcs) - 1:
            doc.add_page_break()
    doc.save(output)


def run_capture(args: argparse.Namespace) -> int:
    try:
        from playwright.sync_api import TimeoutError as PwTimeoutError
        from playwright.sync_api import sync_playwright
    except Exception:
        print(
            "Error: Playwright not installed. Run: pip install playwright && playwright install chromium",
            file=sys.stderr,
        )
        return 2

    out_dir = args.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    wanted_tcs = tc_keys(args.start, args.end)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=args.headless)
        page = browser.new_page(viewport={"width": 1920, "height": 1080})

        try:
            page.goto(args.url, wait_until="domcontentloaded", timeout=args.page_timeout * 1000)
            page.wait_for_selector("button.nav-btn[data-panel='editor']", timeout=30_000)
        except Exception as e:
            browser.close()
            print(f"Error: cannot open GUI at {args.url}: {e}", file=sys.stderr)
            return 2

        for tc in wanted_tcs:
            code = TESTCASES.get(tc)
            if not code:
                print(f"Skipping {tc}: testcase not defined")
                continue

            print(f"[{tc}] running analysis...")
            page.click("button.nav-btn[data-panel='editor']")
            page.wait_for_selector("#panel-editor:not(.hidden)", timeout=10_000)
            page.fill("#ssa-editor", code)
            page.evaluate("useEditorCode()")

            if args.mode in {"classical", "ml"}:
                page.check(f"input[name='mode'][value='{args.mode}']")
            if args.mode == "ml":
                page.evaluate(
                    """(t) => {
                        const s = document.getElementById('sb-threshold');
                        if (!s) return;
                        s.value = String(t);
                        s.dispatchEvent(new Event('input', { bubbles: true }));
                    }""",
                    args.threshold,
                )

            page.click("#run-btn")

            # Wait until analysis ends (DONE or ERROR)
            try:
                page.wait_for_function(
                    """() => {
                        const chip = document.getElementById('status-chip');
                        if (!chip) return false;
                        const t = (chip.textContent || '').toUpperCase();
                        return t.includes('DONE') || t.includes('ERROR');
                    }""",
                    timeout=args.run_timeout * 1000,
                )
            except PwTimeoutError:
                print(f"[{tc}] timed out waiting for analysis completion", file=sys.stderr)
                continue

            status_text = (
                page.locator("#status-chip").inner_text(timeout=5_000).strip().upper()
            )
            if "ERROR" in status_text:
                print(f"[{tc}] analysis error, skipping panel captures", file=sys.stderr)
                continue

            # Allow final paint and rendering (especially CFG SVG)
            time.sleep(args.post_run_wait)

            print(f"[{tc}] capturing 4 GUI panels...")
            for panel_name, panel_selector, suffix in PANEL_MAP:
                page.click(f"button.nav-btn[data-panel='{panel_name}']")
                page.wait_for_selector(f"{panel_selector}:not(.hidden)", timeout=10_000)
                time.sleep(0.25)
                target = out_dir / f"{tc}_{suffix}.png"
                page.locator(panel_selector).screenshot(path=str(target))

        browser.close()

    print(f"Screenshots saved in: {out_dir}")
    return 0


def make_report(images_dir: Path, report_path: Path, start: int, end: int, title: str) -> int:
    images = collect_images(images_dir)
    tcs = tc_keys(start, end)
    ext = report_path.suffix.lower()
    try:
        if ext == ".md":
            render_markdown(report_path, title, tcs, images)
        elif ext == ".docx":
            render_docx(report_path, title, tcs, images)
        else:
            print("Error: report must end with .docx or .md", file=sys.stderr)
            return 2
    except RuntimeError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 2

    found = sum(len(images.get(tc, {})) for tc in tcs)
    expected = len(tcs) * len(PANEL_ORDER)
    print(f"Generated report: {report_path}")
    print(f"Screenshots mapped: {found}/{expected}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Capture GUI-level screenshots for TC01..TC17")
    parser.add_argument("--url", default="http://localhost:3000", help="IDCE GUI URL")
    parser.add_argument("--out-dir", type=Path, required=True, help="Where screenshots will be saved")
    parser.add_argument("--start", type=int, default=1, help="Start testcase number")
    parser.add_argument("--end", type=int, default=17, help="End testcase number")
    parser.add_argument("--mode", choices=["classical", "ml"], default="classical")
    parser.add_argument("--threshold", type=float, default=0.6, help="ML threshold when mode=ml")
    parser.add_argument("--headless", action="store_true", help="Run browser in headless mode")
    parser.add_argument("--run-timeout", type=int, default=180, help="Seconds to wait per testcase run")
    parser.add_argument("--page-timeout", type=int, default=30, help="Seconds to wait for initial page load")
    parser.add_argument("--post-run-wait", type=float, default=0.7, help="Extra seconds before capture")
    parser.add_argument("--report", type=Path, default=None, help="Optional output report (.docx or .md)")
    parser.add_argument("--report-title", default="IDCE Testcase Output Report", help="Title for generated report")
    args = parser.parse_args()

    if args.start < 1 or args.end < args.start:
        print("Error: invalid testcase range", file=sys.stderr)
        return 2

    rc = run_capture(args)
    if rc != 0:
        return rc

    if args.report:
        rep_rc = make_report(args.out_dir, args.report, args.start, args.end, args.report_title)
        if rep_rc != 0:
            return rep_rc

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
